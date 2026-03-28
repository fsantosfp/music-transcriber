import os
import io
import uuid
import re
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, BackgroundTasks
from fastapi.responses import StreamingResponse
from sqlmodel import Session
from app.core.db import engine
from app.models.music import Music, MusicStatus, MusicUpdate
from app.core.config import MAX_UPLOAD_SIZE_MB
import aiofiles
from app.tasks.transcription import process_transcription

from pydantic import BaseModel
from typing import Optional

router = APIRouter()

ALLOWED_EXTENSIONS = {".mp3", ".wav", ".m4a", ".ogg"}
UPLOAD_DIR = "/app/uploads"
if not os.path.exists(UPLOAD_DIR):
    UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
    os.makedirs(UPLOAD_DIR, exist_ok=True)

def get_session():
    with Session(engine) as session:
        yield session

@router.post("/upload", status_code=201, response_model=Music)
async def upload_music_file(background_tasks: BackgroundTasks, file: UploadFile = File(...), session: Session = Depends(get_session)):
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file format. Allowed: {ALLOWED_EXTENSIONS}")

    safe_filename = re.sub(r'[^a-zA-Z0-9_\-.]', '_', file.filename.lower())
    transaction_id = uuid.uuid4()
    new_filename = f"{transaction_id}-{safe_filename}"
    file_path_disk = os.path.join(UPLOAD_DIR, new_filename)
    relative_path = f"uploads/{new_filename}"

    max_size_bytes = MAX_UPLOAD_SIZE_MB * 1024 * 1024
    current_size = 0
    
    try:
        async with aiofiles.open(file_path_disk, 'wb') as out_file:
            while content := await file.read(1024 * 1024):
                current_size += len(content)
                if current_size > max_size_bytes:
                    os.remove(file_path_disk)
                    raise HTTPException(status_code=413, detail="Payload Too Large")
                await out_file.write(content)
    except HTTPException:
        raise
    except Exception as e:
        if os.path.exists(file_path_disk):
            os.remove(file_path_disk)
        raise HTTPException(status_code=500, detail="Failed to save file to disk")

    music_record = Music(
        id=transaction_id,
        filename=safe_filename,
        audio_path=relative_path,
        status=MusicStatus.PENDING
    )
    
    try:
        session.add(music_record)
        session.commit()
        session.refresh(music_record)
        
        # Dispara processo de transcrição em 2o plano
        background_tasks.add_task(process_transcription, str(music_record.id))
        
    except Exception as e:
        if os.path.exists(file_path_disk):
            os.remove(file_path_disk)
        session.rollback()
        raise HTTPException(status_code=500, detail="Failed to persist record in database")

    return music_record

@router.get("/{music_id}", response_model=Music)
def get_music_status(music_id: uuid.UUID, session: Session = Depends(get_session)):
    """
    Retrieve the current status and transcriptions of a specific music upload.
    """
    music = session.get(Music, music_id)
    if not music:
        raise HTTPException(status_code=404, detail="Music not found")
    return music

from fastapi import Query
from pydantic import BaseModel
from typing import Optional

class PaginatedMusicResponse(BaseModel):
    items: list[Music]
    total: int
    page: int
    size: int
    pages: int

@router.get("/", response_model=PaginatedMusicResponse)
def list_all_music(
    session: Session = Depends(get_session),
    page: int = Query(1, ge=1),
    size: int = Query(10, ge=1, le=100),
    q: Optional[str] = None
):
    """
    List uploaded music tracks with pagination and global search.
    """
    from sqlmodel import select, or_, func
    import math

    count_query = select(func.count(Music.id))
    search_query = select(Music)

    if q:
        search_filter = or_(
            Music.filename.contains(q.lower()),
            Music.formatted_transcription.contains(q)
        )
        count_query = count_query.where(search_filter)
        search_query = search_query.where(search_filter)

    total = session.exec(count_query).one()
    search_query = search_query.order_by(Music.created_at.desc()).offset((page - 1) * size).limit(size)
    items = session.exec(search_query).all()

    return {
        "items": items,
        "total": total,
        "page": page,
        "size": size,
        "pages": math.ceil(total / size) if total > 0 else 1
    }

@router.patch("/{music_id}", response_model=Music)
def update_music(music_id: uuid.UUID, music_update: MusicUpdate, session: Session = Depends(get_session)):
    """
    Apply targeted overwrites to properties of an already instantiated upload.
    Mainly utilized for front-end generic saves on JSON Segment Arrays.
    """
    db_music = session.get(Music, music_id)
    if not db_music:
        raise HTTPException(status_code=404, detail="Music not found")
        
    update_data = music_update.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_music, key, value)
        
    # Se o raw_transcription foi atualizado pelo editor, refazemos a formatacao LLM com os dados novos.
    if "raw_transcription" in update_data and update_data["raw_transcription"] is not None:
        from app.services.llm_service import LLMService
        llm_service = LLMService()
        try:
            formatted_text = llm_service.format_transcription(update_data["raw_transcription"])
            db_music.formatted_transcription = formatted_text
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to rebuild formatted transcription from AI: {str(e)}")
        
    session.add(db_music)
    session.commit()
    session.refresh(db_music)
    return db_music

@router.get("/{music_id}/export")
def export_music_lyrics(music_id: uuid.UUID, format: str = "txt", session: Session = Depends(get_session)):
    """
    Generate physical files representing the final transcription.
    """
    music = session.get(Music, music_id)
    if not music:
        raise HTTPException(status_code=404, detail="Music not found")
        
    text = music.formatted_transcription or music.raw_transcription or "Letra não disponível"
    
    if format == "txt":
        buffer = io.BytesIO()
        buffer.write(text.encode("utf-8"))
        buffer.seek(0)
        return StreamingResponse(buffer, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename={music.filename}.txt"})
        
    elif format == "docx":
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx is not installed")
            
        doc = Document()
        doc.add_heading(f"Letra: {music.filename}", 0)
        doc.add_paragraph(text)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={music.filename}.docx"})
        
    elif format == "pdf":
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.utils import simpleSplit
        except ImportError:
            raise HTTPException(status_code=500, detail="reportlab is not installed")
            
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(72, 750, f"Letra: {music.filename}")
        
        c.setFont("Helvetica", 12)
        y = 720
        for line in text.split('\n'):
            line = line.strip('\r')
            if not line.strip():
                y -= 15
                if y < 72:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = 750
                continue
                
            wrapped_lines = simpleSplit(line, "Helvetica", 12, 450)
            for wrapped in wrapped_lines:
                if y < 72:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = 750
                c.drawString(72, y, wrapped)
                y -= 15
        
        c.save()
        buffer.seek(0)
        return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={music.filename}.pdf"})
        
    raise HTTPException(status_code=400, detail="Invalid format requested. Valid formats: txt, docx, pdf.")

@router.post("/{music_id}/retry", response_model=Music)
def retry_music_processing(music_id: uuid.UUID, background_tasks: BackgroundTasks, session: Session = Depends(get_session)):
    """
    Cleans up any previous transcriptions or fallback counters, resetting the pipeline completely
    while aggressively reusing the physical media file uploaded beforehand.
    """
    music = session.get(Music, music_id)
    if not music:
        raise HTTPException(status_code=404, detail="Music not found")
        
    music.status = MusicStatus.PENDING
    music.raw_transcription = None
    music.formatted_transcription = None
    music.vocal_isolation_attempted = False
    
    session.add(music)
    session.commit()
    session.refresh(music)
    
    # Restart the workflow from scratch
    background_tasks.add_task(process_transcription, str(music.id))
    return music

@router.delete("/{music_id}", status_code=204)
def delete_music(music_id: uuid.UUID, session: Session = Depends(get_session)):
    """
    Safely purges the database record alongside all physical media (Original and Vocal stems) from the /uploads folder.
    """
    music = session.get(Music, music_id)
    if not music:
        raise HTTPException(status_code=404, detail="Music not found")
        
    base_dir = "/app"
    
    # 1. Attempt to remove original audio
    if music.audio_path:
        full_audio_path = os.path.join(base_dir, music.audio_path)
        try:
            if os.path.exists(full_audio_path):
                os.remove(full_audio_path)
        except Exception as e:
            print(f"Failed to delete original audio file {full_audio_path}: {e}")
            
    # 2. Attempt to remove isolated vocal traces
    vocal_path = os.path.join(base_dir, "uploads", f"{music.id}_vocals.wav")
    try:
        if os.path.exists(vocal_path):
            os.remove(vocal_path)
    except Exception as e:
        print(f"Failed to delete vocal traces {vocal_path}: {e}")
        
    # 3. Nuke from SQLite
    session.delete(music)
    session.commit()
    return None
